from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/Antonio_Colomba_Modern_Resume.docx"

INK = "202833"
MUTED = "66717F"
BLUE = "174E8F"
TEAL = "0F766E"
LINE = "EEF2F6"


def set_run(run, *, size=10.8, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    for family in ("ascii", "hAnsi", "eastAsia"):
        run._element.rPr.rFonts.set(qn(f"w:{family}"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, *, before=0, after=5, line=1.1, left=0, right=0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)


def add_text(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_run(run, **kwargs)
    return run


def add_bottom_rule(paragraph, color=LINE, size="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def set_cell_width(cell, inches):
    tc_pr = cell._tc.get_or_add_tcPr()
    width = tc_pr.find(qn("w:tcW"))
    if width is None:
        width = OxmlElement("w:tcW")
        tc_pr.append(width)
    width.set(qn("w:w"), str(int(inches * 1440)))
    width.set(qn("w:type"), "dxa")


def add_section(number, title, *, before=14, left=0):
    p = doc.add_paragraph()
    set_para(p, before=before, after=6, line=1.0, left=left)
    add_text(p, f"{number}  ", size=9.0, bold=True, color=TEAL)
    add_text(p, title.upper(), size=10.5, bold=True, color=BLUE)
    return p


def add_role(title, meta="", *, left=0, before=7):
    p = doc.add_paragraph()
    set_para(p, before=before, after=2, line=1.0, left=left)
    add_text(p, title, size=11.0, bold=True, color=INK)
    if meta:
        add_text(p, "   " + meta, size=9.0, color=MUTED)
    return p


def add_bullet(text, *, size=9.55, left=0.18, after=5):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, before=0, after=after, line=1.08, left=left)
    add_text(p, text, size=size, color=INK)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.44)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(10.1)
styles["Normal"].paragraph_format.space_after = Pt(4)
styles["Normal"].paragraph_format.line_spacing = 1.08
styles["List Bullet"].font.name = "Aptos"
styles["List Bullet"].font.size = Pt(9.55)
styles["List Bullet"].paragraph_format.left_indent = Inches(0.3)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.14)
styles["List Bullet"].paragraph_format.space_after = Pt(3)
styles["List Bullet"].paragraph_format.line_spacing = 1.08

top = doc.add_table(rows=1, cols=2)
top.autofit = False
left, right = top.rows[0].cells
set_cell_width(left, 4.95)
set_cell_width(right, 1.85)
for cell in top.rows[0].cells:
    clear_cell_borders(cell)

p = left.paragraphs[0]
set_para(p, after=0, line=0.9)
add_text(p, "ANTONIO", size=35, bold=True, color=INK)
p = left.add_paragraph()
set_para(p, before=0, after=2, line=0.9)
add_text(p, "COLOMBA", size=35, bold=True, color=INK)
p = left.add_paragraph()
set_para(p, before=0, after=0, line=0.95)
add_text(p, "Creative Technology & AI Workflow Operator", size=13.4, bold=True, color=BLUE)

p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_para(p, before=8, after=2, line=1.0)
add_text(p, "NYC / Remote", size=9.5, bold=True, color=INK)
for text in ["(561) 271-1163", "getantonio@gmail.com", "project walkthroughs available"]:
    p = right.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_para(p, after=2, line=1.0)
    add_text(p, text, size=8.8, color=TEAL if "project" in text else MUTED)

summary = doc.add_paragraph()
set_para(summary, before=24, after=14, line=1.2, left=0.34, right=0.16)
add_text(
    summary,
    "High-end retouching, VFX, animation, compositing, and agency/client production foundation now pointed at AI-assisted technical work. "
    "Not selling current senior 3D speed. Selling visual judgment, quality control, tool learning, client delivery, organization, and follow-through.",
    size=11.0,
)

fit = doc.add_paragraph()
set_para(fit, before=0, after=22, line=1.1, left=1.0, right=0.18)
add_text(fit, "Best bought for: ", size=9.65, bold=True, color=TEAL)
add_text(
    fit,
    "production coordination, creative operations, AI workflow support, agent operations, technical support, junior project support.",
    size=9.65,
    color=MUTED,
)
add_bottom_rule(fit, size="1")

add_section("01", "Current technical project work", before=10, left=0.24)
add_role("Token Factory / TokenHub", "local/private crypto product prototype", left=0.24)
add_bullet("Built token creation and management flows covering configuration, supply, transfer restrictions, wallet addresses, security settings, liquidity concepts, and post-creation steps.", left=0.4)
add_bullet("Explored multi-network token creation concepts across Ethereum Sepolia, Arbitrum Sepolia, Optimism Sepolia, and Polygon Amoy.", left=0.4)

add_role("Trading & Crypto Tooling", "AI-assisted local projects", left=0.04, before=9)
add_bullet("Worked on GMX tooling, signal dashboards, strategy systems, arbitrage prototypes, and DLtrade-style interfaces.", left=0.2)
add_bullet("Used AI-assisted development to read documentation, debug issues, organize workflows, and turn product ideas into local prototypes.", left=0.2)

add_role("AI Workflow Practice", "current direction", left=0.48, before=9)
add_bullet("Organizing messy product ideas into tasks, project notes, testable flows, short write-ups, and walkthrough-ready proof assets.", left=0.64)

note = doc.add_paragraph()
set_para(note, before=7, after=15, line=1.1, left=0.82, right=0.18)
add_text(note, "Presentation plan: ", size=9.05, bold=True, color=TEAL)
add_text(note, "local/private projects should be shown through screenshots, walkthroughs, and short project summaries, not live product claims.", size=9.05, color=MUTED)

tools = doc.add_paragraph()
set_para(tools, before=7, after=16, line=1.1, left=0.08, right=0.15)
add_text(tools, "Tools: ", size=9.05, bold=True, color=TEAL)
add_text(
    tools,
    "Photoshop, compositing, color, Nuke, After Effects, Premiere, InDesign, Maya, Houdini, Python, MEL, Mac, Windows, Linux, troubleshooting, documentation, AI-assisted workflows, Next.js, TypeScript, crypto, payments.",
    size=8.85,
    color=MUTED,
)

divider = doc.add_paragraph()
set_para(divider, before=13, after=9, line=1.0)
add_text(divider, "ANTONIO COLOMBA", size=9.6, bold=True, color=INK)
add_text(divider, "   background and education", size=8.8, color=MUTED)
add_bottom_rule(divider, size="1")

add_section("02", "Professional foundation", before=8)
add_role("Freelance Engagements, High-End Digital Retoucher", "NYC | dates to confirm")
add_bullet("Delivered high-end retouching and visual finishing for agency, studio, and commercial clients including Foam Digital, McCann Erickson, Urban Studio, Nucleus Imaging, Digital Evolution, EMR Systems, Fuel Digital, Color Edge, and Young & Rubicam Inc.")
add_bullet("Applied image correction, color matching, compositing, proofing, and client-direction revisions.")

add_role("291 Digital, High-End Digital Retoucher", "NYC | 12/2005 - 09/2006", left=0.32, before=8)
add_bullet("Performed color correction and proof matching for high-end retouching projects requiring precision and consistency.", left=0.48)

add_role("FCB World Wide, High-End Digital Retoucher", "NYC | 09/2000 - 09/2001", before=8)
add_bullet("Produced high-resolution art, visual effects, retouching, color correction, and creative compositions.")

add_section("03", "Education", before=10, left=0.22)
for item in [
    "Digital Media Arts College, M.F.A. Visual Effects Animation, Highest Academic Distinction.",
    "Digital Media Arts College, B.F.A. 3D Character Animation, Summa Cum Laude, First in Class 2015.",
    "CG Spectrum / RentAmentor, 3D Character Animation; Renaissance Center / 3D Buzz, 3D Animation; School of Visual Arts, Commercial Photography.",
]:
    add_bullet(item, size=9.0, left=0.4, after=3)

doc.save(OUT)
print(OUT)
