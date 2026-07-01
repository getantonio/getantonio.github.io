from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from shutil import copyfile


OUT = "docs/Antonio_Colomba_Operations_Support_Resume.docx"
SITE_OUT = "Antonio_Colomba_Operations_Support_Resume.docx"

INK = "1F2933"
MUTED = "5B6673"
BLUE = "164E8F"
TEAL = "0F766E"
RULE = "E8EDF2"


def set_run(run, *, size=10, bold=False, color=INK, italic=False):
    run.font.name = "Aptos"
    for family in ("ascii", "hAnsi", "eastAsia"):
        run._element.rPr.rFonts.set(qn(f"w:{family}"), "Aptos")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(p, *, before=0, after=5, line=1.08, left=0, right=0):
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    fmt.left_indent = Inches(left)
    fmt.right_indent = Inches(right)


def add_text(p, text, **kwargs):
    run = p.add_run(text)
    set_run(run, **kwargs)
    return run


def add_rule(p):
    p_pr = p._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "1")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), RULE)


def section_head(number, title, *, before=11):
    p = doc.add_paragraph()
    set_para(p, before=before, after=5, line=1.0)
    add_text(p, f"{number}  ", size=8.6, bold=True, color=TEAL)
    add_text(p, title.upper(), size=10.1, bold=True, color=BLUE)
    return p


def role(title, meta="", *, before=7, left=0):
    p = doc.add_paragraph()
    set_para(p, before=before, after=1, line=1.0, left=left)
    add_text(p, title, size=10.6, bold=True)
    if meta:
        add_text(p, "   " + meta, size=8.7, color=MUTED)
    return p


def bullet(text, *, left=0.18, after=3, size=9.2):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, before=0, after=after, line=1.08, left=left)
    add_text(p, text, size=size)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.68)
section.bottom_margin = Inches(0.46)
section.left_margin = Inches(0.62)
section.right_margin = Inches(0.62)

styles = doc.styles
styles["Normal"].font.name = "Aptos"
styles["Normal"].font.size = Pt(9.6)
styles["Normal"].paragraph_format.space_after = Pt(4)
styles["Normal"].paragraph_format.line_spacing = 1.08
styles["List Bullet"].font.name = "Aptos"
styles["List Bullet"].font.size = Pt(9.2)
styles["List Bullet"].paragraph_format.left_indent = Inches(0.3)
styles["List Bullet"].paragraph_format.first_line_indent = Inches(-0.14)
styles["List Bullet"].paragraph_format.space_after = Pt(3)
styles["List Bullet"].paragraph_format.line_spacing = 1.08

name = doc.add_paragraph()
set_para(name, after=0, line=0.92)
add_text(name, "ANTONIO COLOMBA", size=27, bold=True)

contact = doc.add_paragraph()
set_para(contact, after=3, line=1.0)
add_text(contact, "NYC / Remote  |  (561) 271-1163  |  getantonio@gmail.com", size=9.2, color=MUTED)

title = doc.add_paragraph()
set_para(title, after=9, line=1.0)
add_text(title, "Operations Support | Production Coordination | Customer Support", size=12.2, bold=True, color=BLUE)
add_rule(title)

summary = doc.add_paragraph()
set_para(summary, before=9, after=9, line=1.12)
add_text(
    summary,
    "Reliable, detail-oriented worker with a background in deadline-driven creative service production environments, client revisions, visual quality control, troubleshooting, AI content creation workflows, and agentic engineered technical projects. Strong fit for support, AI onboarding and adaptation of operations, admin, production coordination, temp, and technical support roles.",
    size=9.8,
)

fit = doc.add_paragraph()
set_para(fit, before=0, after=9, line=1.08, left=0.2)
add_text(fit, "Best fit roles: ", size=9.0, bold=True, color=TEAL)
add_text(
    fit,
    "Customer Support Representative, Operations Assistant, Data Entry Clerk, Office Assistant, Production Coordinator, Technical Support Associate, Customer Success Coordinator, Onboarding Assistant, Temp Office Support.",
    size=8.9,
    color=MUTED,
)

section_head("01", "Core strengths", before=7)
for item in [
    "Calm written communication, organized follow-up, and careful handling of details.",
    "Client/project background with revisions, deadline pressure, quality checks, proofing, and production standards.",
    "Comfortable learning software tools, troubleshooting issues, documenting process, and following structured workflows.",
    "Ready to support ticket handling, research, CRM updates, data entry, checklists, file organization, customer communication, and AI content workflows.",
]:
    bullet(item)

section_head("02", "Professional experience", before=10)
role("Freelance Engagements, High-End Digital Retoucher", "NYC | dates to confirm")
bullet("Delivered detailed visual finishing and retouching work for agency, studio, and commercial clients.")
bullet("Worked with teams including Foam Digital, McCann Erickson, Urban Studio, Nucleus Imaging, Digital Evolution, EMR Systems, Fuel Digital, Color Edge, and Young & Rubicam Inc.")
bullet("Managed image correction, color matching, compositing, proofing, revision cycles, client direction, and quality-control expectations.")

role("291 Digital, High-End Digital Retoucher", "NYC | 12/2005 - 09/2006", left=0.28)
bullet("Performed color correction and proof matching for high-end retouching projects requiring precision and consistency.", left=0.44)

role("FCB World Wide, High-End Digital Retoucher", "NYC | 09/2000 - 09/2001")
bullet("Produced high-resolution art, visual effects, retouching, color correction, and creative compositions.")

section_head("03", "Current project work", before=10)
role("Agentic Engineered Technical Projects", "local/private")
bullet("Built and organized local crypto, token, trading, AI content creation, and workflow projects with AI-assisted tools.")
bullet("Practiced troubleshooting, reading documentation, organizing tasks, testing flows, and turning messy ideas into usable workflows.")

section_head("04", "Work readiness", before=10)
ready = doc.add_paragraph()
set_para(ready, before=0, after=7, line=1.08, left=0.16)
add_text(
    ready,
    "Available for entry-level support, operations, admin, production coordination, customer success, technical support, and temp assignments. Open to in-person, hybrid, remote, temporary, contract, part-time, or full-time work.",
    size=9.1,
)

tools = doc.add_paragraph()
set_para(tools, before=2, after=7, line=1.08)
add_text(tools, "Tools: ", size=8.8, bold=True, color=TEAL)
add_text(
    tools,
    "Photoshop, InDesign, Premiere, After Effects, Nuke, Maya, Houdini, Mac, Windows, Linux, troubleshooting, documentation, spreadsheets, AI-assisted workflows, Higgsfield, generative media workflows, Next.js, TypeScript, Python basics, crypto/product research.",
    size=8.65,
    color=MUTED,
)

section_head("05", "Education", before=10)
schools = [
    ("Digital Media Arts College", "South Florida", "M.F.A. Visual Effects Animation, Highest Academic Distinction."),
    ("Digital Media Arts College", "South Florida", "B.F.A. 3D Character Animation, Summa Cum Laude, First in Class 2015."),
    ("CG Spectrum / Rent-A-Mentor", "Remote", "3D Character Animation."),
    ("Renaissance Center / 3D Buzz", "Tennessee", "3D Animation."),
    ("School of Visual Arts", "NYC", "Commercial Photography."),
]
for school, location, detail in schools:
    edu = doc.add_paragraph()
    set_para(edu, before=1, after=0, line=1.0, left=0.16)
    add_text(edu, school, size=8.85, bold=True)
    add_text(edu, f"   {location}", size=8.35, color=MUTED)
    edu_detail = doc.add_paragraph()
    set_para(edu_detail, before=0, after=2, line=1.0, left=0.32)
    add_text(edu_detail, detail, size=8.55, color=MUTED)

doc.save(OUT)
copyfile(OUT, SITE_OUT)
print(OUT)
print(SITE_OUT)
