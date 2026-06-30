from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/Antonio_Ripple_Fintech_Resume_Draft.docx"


def set_run(run, *, size=10.5, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, *, before=0, after=4, line=1.0):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=3, line=1.0)
    r = p.add_run(text.upper())
    set_run(r, size=10.5, bold=True, color="1F4D78")
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph(p, before=0, after=2, line=1.05)
    r = p.add_run(text)
    set_run(r, size=9.7)
    return p


def add_role(doc, company, title, dates):
    p = doc.add_paragraph()
    set_paragraph(p, before=4, after=1, line=1.0)
    r = p.add_run(f"{company} | {title}")
    set_run(r, size=10.3, bold=True)
    r = p.add_run(f"    {dates}")
    set_run(r, size=9.5, color="555555")
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.55)
section.bottom_margin = Inches(0.55)
section.left_margin = Inches(0.65)
section.right_margin = Inches(0.65)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10)
normal.paragraph_format.space_after = Pt(3)
normal.paragraph_format.line_spacing = 1.05

for style_name in ["List Bullet", "List Number"]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(9.7)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.line_spacing = 1.05

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph(title, before=0, after=1, line=1.0)
r = title.add_run("ANTONIO COLOMBA")
set_run(r, size=18, bold=True, color="0B2545")

contact = doc.add_paragraph()
contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph(contact, before=0, after=4, line=1.0)
r = contact.add_run("[City, State] | [Phone] | [Email] | [LinkedIn] | [GitHub / Portfolio]")
set_run(r, size=9.4, color="555555")

headline = doc.add_paragraph()
headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_paragraph(headline, before=0, after=6, line=1.0)
r = headline.add_run("Creative Technology & AI Workflow Operator | Production, Support, Crypto Projects")
set_run(r, size=10.3, bold=True, color="1F4D78")

add_heading(doc, "Profile")
p = doc.add_paragraph()
set_paragraph(p, before=0, after=4, line=1.05)
r = p.add_run(
    "Creative technology operator with a foundation in high-end photo retouching, visual effects, animation, compositing, and "
    "production-quality client work. Currently focused on AI-assisted development, workflow organization, project coordination, "
    "support, operations, and fintech/crypto product learning. Not positioning as a current senior hands-on 3D/VFX artist; "
    "strongest value is visual/technical judgment, production discipline, tool learning, documentation, and making messy workflows usable."
)
set_run(r, size=9.8)

add_heading(doc, "Target Roles")
p = doc.add_paragraph()
set_paragraph(p, before=0, after=4, line=1.05)
r = p.add_run(
    "Production Coordinator | Creative Operations Assistant | AI Workflow Coordinator | Agent Operations Assistant | Technical Support | "
    "Customer Support | Operations Associate | Implementation Coordinator | Junior Project Coordinator | Product Operations Assistant"
)
set_run(r, size=9.6)

add_heading(doc, "Core Strengths")
add_bullet(doc, "Visual production foundation: retouching, color, compositing, animation, and image quality judgment.")
add_bullet(doc, "Client/project work across agencies and creative teams, including deadline-driven delivery.")
add_bullet(doc, "Strong fit for coordination, operations, documentation, QA, research, customer-facing, and AI workflow support.")
add_bullet(doc, "Comfortable working with AI-assisted coding tools to prototype ideas, test workflows, and organize technical projects.")
add_bullet(doc, "Hands-on web3 project exposure through Token Factory / TokenHub, including token creation and exchange/listing concepts.")

add_heading(doc, "Professional Experience")
add_role(doc, "Freelance Engagements", "High-End Digital Retoucher", "NYC | Dates to confirm")
add_bullet(doc, "Delivered high-end digital retouching and visual finishing work for agency, studio, and commercial clients.")
add_bullet(doc, "Worked with brands and teams including Foam Digital, McCann Erickson, Urban Studio, Nucleus Imaging, Digital Evolution, EMR Systems, Fuel Digital, Color Edge, and Young & Rubicam Inc.")
add_bullet(doc, "Applied detailed image correction, color matching, compositing, and proofing to meet professional production standards.")
add_bullet(doc, "Collaborated with creative stakeholders and adapted work to client direction, quality expectations, and tight deadlines.")

add_role(doc, "291 Digital", "High-End Digital Retoucher", "NYC | 12/2005 - 09/2006")
add_bullet(doc, "Performed meticulous color correcting and proof matching for high-end retouching projects.")
add_bullet(doc, "Supported prestigious client work requiring precision, consistency, and visual quality control.")

add_role(doc, "FCB World Wide", "High-End Digital Retoucher", "NYC | 09/2000 - 09/2001")
add_bullet(doc, "Produced high-resolution art, visual effects, retouching, color correction, and creative compositions.")
add_bullet(doc, "Supported commercial creative production with detail-focused image finishing.")

add_heading(doc, "Web3 / Fintech Project Experience")
add_role(doc, "Token Factory & Exchange", "Crypto / Fintech Product Project", "Local prototype")
add_bullet(doc, "Built a hands-on crypto product prototype focused on custom token creation, token management, and exchange/listing concepts.")
add_bullet(doc, "Worked through token product flows including token configuration, supply, transfer restrictions, wallet addresses, security settings, liquidity concepts, and post-creation steps.")
add_bullet(doc, "Explored multi-network token creation concepts across Ethereum Sepolia, Arbitrum Sepolia, Optimism Sepolia, and Polygon Amoy.")
add_bullet(doc, "Current status: not publicly accessible. Use screenshots, local code, and a short walkthrough write-up as proof instead of a live website link.")

add_role(doc, "Trading & Crypto Tooling Projects", "AI-Assisted Development", "Local/private projects")
add_bullet(doc, "Built and iterated on local trading and crypto applications including GMX tooling, signal dashboards, trading strategy systems, and arbitrage prototypes.")
add_bullet(doc, "Worked with project structures using Next.js, TypeScript, Python, wallet connections, market data, dashboards, and strategy/backtesting concepts.")
add_bullet(doc, "Used AI-assisted development to explore product workflows, debug issues, read documentation, and turn ideas into working local prototypes.")
add_bullet(doc, "Current status: local/private projects; present through screenshots, repository walk-throughs, and short project summaries rather than public production claims.")

add_role(doc, "Fintech Company Research Brief", "Payments / Crypto / Compliance", "")
add_bullet(doc, "Compare Ripple, Stripe, Wise, Coinbase, and Circle by customer, business model, product, payments, compliance, custody, liquidity, and blockchain usage.")
add_bullet(doc, "Create a one-page glossary covering KYC, AML, sanctions, custody, settlement, liquidity, stablecoins, on/off-ramps, and cross-border payments.")

add_heading(doc, "Work Readiness")
add_bullet(doc, "Available for production coordination, creative operations, AI workflow support, agent operations, entry-level support, customer success, research, and coordination roles.")
add_bullet(doc, "Comfortable starting with ticket handling, documentation, data entry, research, customer communication, CRM updates, and process checklists.")
add_bullet(doc, "Able to discuss professional creative production experience, AI-assisted project work, and self-directed crypto/trading prototypes.")
add_bullet(doc, "No formal fintech employment yet.")

add_heading(doc, "Education")
add_bullet(doc, "Digital Media Arts College, M.F.A. Visual Effects Animation, Highest Academic Distinction.")
add_bullet(doc, "Digital Media Arts College, B.F.A. 3D Character Animation, Summa Cum Laude, First in Class 2015.")
add_bullet(doc, "CG Spectrum / RentAmentor, 3D Character Animation; Renaissance Center / 3D Buzz, 3D Animation; School of Visual Arts, Commercial Photography.")

add_heading(doc, "Keyword Bank - Use Only If True")
p = doc.add_paragraph()
set_paragraph(p, before=0, after=0, line=1.05)
r = p.add_run(
    "Photo retouching, Photoshop, color matching, compositing, rotoscoping, 3D tracking, set extension, character animation, video "
    "editing, Unreal Engine, Nuke, Premiere, After Effects, InDesign, Maya, Houdini, 3DEqualizer, Boujou, Vex, Python, MEL scripting, "
    "Mac OS, Windows, Linux, troubleshooting, customer support, operations, documentation, research, payments, fintech, blockchain, crypto."
)
set_run(r, size=8.8, color="555555")

doc.save(OUT)
print(OUT)
